import streamlit as st
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.summarize import load_summarize_chain
from docx import Document as DocxDocument
from docx.shared import Pt

# Streamlit interface
st.title("Document Summary Generator")
uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")

if uploaded_file is not None:
    pdf = PdfReader(uploaded_file)

    # Read text from PDF
    text = ''
    for page in pdf.pages:
        content = page.extract_text()
        if content:
            text += content + "\n"

    docs = [Document(page_content=text)]

    llm = ChatGroq(groq_api_key='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', model_name='llama3-70b-8192', temperature=0.2, top_p=0.2)

    template = '''Write a very concise, well-explained, point-wise, short summary of the following text. Provide good and user-acceptable response.
    '{text}
    Create section-wise summary.
    The format should be:
    Overview
    Involved Parties
    Key Events
    Key Findings
    The summary should be formatted correctly as per this
    The Key events section needs to be more elaborate and each event should have a heading to it also in a pointer manner'
    '''

    prompt = PromptTemplate(input_variables=['text'], template=template)
    
    chain = load_summarize_chain(llm, chain_type='stuff', prompt=prompt, verbose=False)
    output_summary = chain.invoke(docs)
    output = output_summary['output_text']

    st.write("### Summary:")
    st.write(output)

    # Create a DOCX file
    doc = DocxDocument()

    sections = ["Overview", "Involved Parties", "Key Events", "Key Findings"]
    formatted_summary = {}

    for i in range(len(sections)):
        start = output.find(sections[i])
        end = output.find(sections[i + 1]) if i + 1 < len(sections) else len(output)
        if start != -1:
            formatted_summary[sections[i]] = output[start + len(sections[i]):end].strip()

    for section, content in formatted_summary.items():
        doc.add_paragraph(section, style='Heading 1')  # Bold header
        paragraph = doc.add_paragraph()
        paragraph.add_run(content).font.size = Pt(11)  # Regular text

    # Save the DOCX file
    doc_output_path = "summary_output.docx"
    doc.save(doc_output_path)

    # Provide download link
    with open(doc_output_path, "rb") as doc_file:
        st.download_button("Download Summary DOCX", doc_file, file_name="summary_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.write("Please upload a PDF file.")
